from dataclasses import dataclass, field
from pathlib import Path
import re


@dataclass
class Document:
    content: str
    metadata: dict = field(default_factory=dict)
    doc_id: str = ""


@dataclass
class Chunk:
    content: str
    metadata: dict = field(default_factory=dict)
    chunk_id: str = ""


class DocumentLoader:
    """Loads documents from text, PDF, DOCX, or raw strings."""

    def load_text_file(self, path: str) -> Document:
        p = Path(path)
        content = p.read_text(encoding="utf-8")
        return Document(
            content=content,
            metadata={"source": str(p), "filename": p.name, "type": "text"},
            doc_id=p.stem,
        )

    def load_pdf(self, path: str) -> Document:
        try:
            import PyPDF2
        except ImportError:
            raise ImportError("Install PyPDF2: pip install PyPDF2")

        p = Path(path)
        pages = []
        with open(p, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                pages.append(page.extract_text() or "")
        content = "\n\n".join(pages)
        return Document(
            content=content,
            metadata={"source": str(p), "filename": p.name, "type": "pdf", "pages": len(pages)},
            doc_id=p.stem,
        )

    def load_docx(self, path: str) -> Document:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")

        p = Path(path)
        doc = DocxDocument(str(p))
        content = "\n".join(para.text for para in doc.paragraphs)
        return Document(
            content=content,
            metadata={"source": str(p), "filename": p.name, "type": "docx"},
            doc_id=p.stem,
        )

    def load_from_string(self, text: str, doc_id: str = "manual", metadata: dict = None) -> Document:
        return Document(
            content=text,
            metadata=metadata or {"source": "manual", "type": "text"},
            doc_id=doc_id,
        )

    def load_directory(self, directory: str) -> list[Document]:
        """Load all supported files from a directory."""
        docs = []
        p = Path(directory)
        loaders = {".txt": self.load_text_file, ".pdf": self.load_pdf, ".docx": self.load_docx}
        for file in p.iterdir():
            loader = loaders.get(file.suffix.lower())
            if loader:
                docs.append(loader(str(file)))
        return docs


class TextChunker:
    """Splits documents into overlapping chunks for embedding."""

    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 64):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk_document(self, doc: Document) -> list[Chunk]:
        text = doc.content.strip()
        if not text:
            return []

        sentences = self._split_into_sentences(text)
        chunks = []
        current_chunk = []
        current_length = 0
        chunk_index = 0

        for sentence in sentences:
            sentence_len = len(sentence)
            if current_length + sentence_len > self.chunk_size and current_chunk:
                chunk_text = " ".join(current_chunk)
                chunks.append(
                    Chunk(
                        content=chunk_text,
                        metadata={**doc.metadata, "doc_id": doc.doc_id, "chunk_index": chunk_index},
                        chunk_id=f"{doc.doc_id}_chunk_{chunk_index}",
                    )
                )
                chunk_index += 1
                # Keep overlap: retain last few sentences
                overlap_text = " ".join(current_chunk)
                overlap_words = overlap_text.split()[-self.chunk_overlap:]
                current_chunk = [" ".join(overlap_words)]
                current_length = len(" ".join(overlap_words))

            current_chunk.append(sentence)
            current_length += sentence_len + 1

        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunks.append(
                Chunk(
                    content=chunk_text,
                    metadata={**doc.metadata, "doc_id": doc.doc_id, "chunk_index": chunk_index},
                    chunk_id=f"{doc.doc_id}_chunk_{chunk_index}",
                )
            )

        return chunks

    def chunk_documents(self, docs: list[Document]) -> list[Chunk]:
        all_chunks = []
        for doc in docs:
            all_chunks.extend(self.chunk_document(doc))
        return all_chunks

    def _split_into_sentences(self, text: str) -> list[str]:
        text = re.sub(r"\s+", " ", text)
        sentences = re.split(r"(?<=[.!?])\s+", text)
        return [s.strip() for s in sentences if s.strip()]
